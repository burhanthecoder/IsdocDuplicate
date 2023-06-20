from fastapi import FastAPI, File, UploadFile,HTTPException
import pdfplumber
from docx import Document
from utils import checkDocumentDuplicate
import  os
app = FastAPI()

@app.get("/")
async def root():
    return {"message": "Hello World"}



@app.post("/uploadfile/")
async def create_upload_file(file: UploadFile):
    data = ''
    contents = await file.read()
    with open(file.filename, "wb") as f:
        f.write(contents)
    if file.filename.endswith(".pdf"):
        with pdfplumber.open(file.filename) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                data = data + text
    elif file.filename.endswith((".docx", ".doc")):
        doc = Document(file.filename)
        for paragraph in doc.paragraphs:
            print(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    data = data + cell.text
    else:
        raise Exception("File format not supported")
    isDuplicate = checkDocumentDuplicate(data)
    os.remove(file.filename)
    return { "isDocDuplicate": isDuplicate}
